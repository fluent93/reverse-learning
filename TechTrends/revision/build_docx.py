"""Convert Manuscript_v2_draft.md to a highlighted, APA-formatted .docx.

Markdown conventions used by the draft:
  ==text==          new/revised text  -> yellow highlight
  *text*            italic (used in references)
  # / ## / ###      heading levels
  [TABLE]...[/TABLE] pipe-separated table, first row is header
  [FIGURE 1 HERE]   insert the framework figure
"""
import argparse
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_LINE_SPACING,
    WD_COLOR_INDEX,
    WD_TAB_ALIGNMENT,
)
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/home/ubuntu/work/reverse-learning/TechTrends/revision"
parser = argparse.ArgumentParser()
parser.add_argument("--src", default=f"{BASE}/Manuscript_v2_draft.md")
parser.add_argument("--out", default=f"{BASE}/Main_Manuscript_Reverse_Learning_Framework_R1.docx")
parser.add_argument("--fig", default=f"{BASE}/Figure1_Reverse_Learning_Framework.png")
args = parser.parse_args()
SRC = args.src
OUT = args.out
FIG = args.fig

doc = Document()
doc.core_properties.author = ""
doc.core_properties.last_modified_by = ""
doc.core_properties.comments = ""
doc.core_properties.title = "From AI-Generated Output to Learner Ownership"

section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.45)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)


header = section.header.paragraphs[0]
header.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
header_run = header.add_run("Running head: REVERSE LEARNING FRAMEWORK")
header_run.font.name = "Times New Roman"
header_run.font.size = Pt(10)
header.add_run("\t")
add_page_field(header)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

for lvl in ("Heading 1", "Heading 2", "Heading 3"):
    h = doc.styles[lvl]
    h.font.name = "Times New Roman"
    h.font.bold = True
    h.font.color.rgb = None
    h.font.size = Pt(14 if lvl == "Heading 1" else 12)

doc.styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.styles["Heading 2"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.styles["Heading 3"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_runs(paragraph, text, force_hl=False):
    """Add runs, toggling highlight on == and italics on *...*."""
    hl = force_hl
    for part in re.split(r"(==)", text):
        if part == "==":
            hl = not hl
            continue
        if not part:
            continue
        for seg in re.split(r"(\*[^*\n]+\*)", part):
            if not seg:
                continue
            italic = seg.startswith("*") and seg.endswith("*") and len(seg) > 2
            run = paragraph.add_run(seg[1:-1] if italic else seg)
            run.italic = italic or None
            if hl or force_hl:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def set_apa_borders(table):
    """Horizontal rules only: top, header-bottom, table-bottom."""
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    for edge in ("left", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    tblPr.append(borders)
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcB = OxmlElement("w:tcBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "000000")
        tcB.append(bottom)
        tcPr.append(tcB)


with open(SRC) as f:
    blocks = [b.strip() for b in f.read().split("\n\n") if b.strip()]

in_references = False
first_heading = True
abstract_next = False
after_table_or_figure_number = False

i = 0
while i < len(blocks):
    block = blocks[i]

    # figure placeholder
    if block == "[FIGURE 1 HERE]":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(FIG, width=Inches(6.3))
        i += 1
        continue

    # APA table/figure number and title. Number is bold on its own line;
    # the title is italic on the following line.
    plain_block = block.replace("==", "")
    label_match = re.match(r"^(Table|Figure)\s+(\d+)(?:\.\s+(.+))?$", plain_block)
    if label_match and "\n" not in block:
        kind, number, inline_title = label_match.groups()
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_after = Pt(0)
        if inline_title:
            marker_match = re.match(rf"^{kind}\s+(==)?{number}(==)?", block)
            number_text = f"{kind} " + (f"=={number}==" if marker_match and marker_match.group(1) else number)
        else:
            number_text = block
        add_runs(p, number_text)
        for run in p.runs:
            run.bold = True
        if inline_title:
            title_p = doc.add_paragraph()
            title_p.paragraph_format.first_line_indent = Inches(0)
            title_p.paragraph_format.space_after = Pt(6)
            add_runs(title_p, inline_title)
            for run in title_p.runs:
                run.italic = True
            after_table_or_figure_number = False
        else:
            after_table_or_figure_number = True
        i += 1
        continue

    # table: starts with [TABLE]; rows may share the block or span blocks
    if block.strip("=").startswith("[TABLE]"):
        rows = []
        done = False
        while i < len(blocks) and not done:
            for line in blocks[i].split("\n"):
                line = line.strip()
                core = line.strip("=")
                if not line or core == "[TABLE]":
                    continue
                if core == "[/TABLE]":
                    done = True
                    break
                hl = line.startswith("==") and line.endswith("==")
                cells = [c.strip() for c in core.split("|")]
                rows.append((cells, hl))
            i += 1
        ncols = max(len(r[0]) for r in rows)
        table = doc.add_table(rows=len(rows), cols=ncols)
        table.autofit = True
        for r, (cells, hl) in enumerate(rows):
            for c in range(ncols):
                cell = table.cell(r, c)
                cell.paragraphs[0].paragraph_format.space_after = Pt(2)
                text = cells[c] if c < len(cells) else ""
                add_runs(cell.paragraphs[0], text, force_hl=hl)
                if r == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
        set_apa_borders(table)
        doc.add_paragraph()
        continue

    # headings
    m = re.match(r"^(=*)(#{1,3})\s+(.*?)(=*)$", block, re.S)
    if m and "\n" not in block:
        marks, hashes, text, _ = m.groups()
        hl = bool(marks)
        if first_heading:
            # manuscript title
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_runs(p, text)
            for run in p.runs:
                run.bold = True
            first_heading = False
        else:
            if text.strip("=").startswith("References"):
                doc.add_page_break()
            p = doc.add_heading("", level=len(hashes))
            add_runs(p, text.strip("="), force_hl=hl)
            for run in p.runs:
                run.font.name = "Times New Roman"
            if text.strip().startswith("References"):
                in_references = True
            abstract_next = text.strip("=").startswith("Abstract")
        i += 1
        continue

    # normal paragraph (may span multiple lines)
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(0)
    if in_references:
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.5)
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    else:
        if abstract_next or block.startswith("Keywords:") or after_table_or_figure_number or block.startswith(("*Note.*", "==*Note.*")):
            pf.first_line_indent = Inches(0)
        else:
            pf.first_line_indent = Inches(0.5)
    add_runs(p, " ".join(line.strip() for line in block.split("\n")))
    if after_table_or_figure_number:
        for run in p.runs:
            run.italic = True
        pf.space_after = Pt(6)
        after_table_or_figure_number = False
    abstract_next = False
    i += 1

doc.save(OUT)
print("saved", OUT)
