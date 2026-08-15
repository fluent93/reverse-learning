"""Build an improved TechTrends title page while preserving author placeholders."""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT = "/home/ubuntu/work/reverse-learning/TechTrends/revision/Title_Page_Reverse_Learning_Framework_R2.docx"
TITLE = "From AI-Generated Output to Learner Ownership: A Reverse Learning Framework for Generative AI-Mediated Education"

doc = Document()
doc.core_properties.author = "Changhan Ryu"
doc.core_properties.last_modified_by = ""
doc.core_properties.comments = ""
doc.core_properties.title = TITLE
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
style.paragraph_format.space_after = Pt(8)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(TITLE)
run.bold = True
run.font.size = Pt(14)


def add_field(label, value, missing=False):
    p = doc.add_paragraph()
    lead = p.add_run(f"{label}: ")
    lead.bold = True
    value_run = p.add_run(value)
    if missing:
        value_run.font.highlight_color = WD_COLOR_INDEX.YELLOW


add_field("Author", "Changhan Ryu")
add_field("Affiliations", "Arizona State University; LG Electronics")
add_field("Corresponding author", "Changhan Ryu")
add_field("Email", "fluent93@gmail.com")
add_field("Phone", "+82-10-3078-9744")
add_field("Postal address", "205-21, Magokseo-ro, Gangseo-gu, Seoul, Korea 07593")
add_field("ORCID", "https://orcid.org/0009-0000-0817-5743")
add_field(
    "Keywords",
    "academic integrity; AI literacy; explainable ownership; generative AI; human-AI interaction; instructional design; learner ownership; learning design; metacognition; Reverse Learning",
)
add_field("Manuscript word count", "Approximately 5,200 words, including the abstract and excluding references, tables, and the figure")
add_field("Tables and figures", "4 tables; 1 figure")
add_field("Acknowledgments", "None declared")
add_field("Funding", "No external funding was received for this manuscript")
add_field("Conflict of interest", "The author declares no conflict of interest")
add_field(
    "Author biographical statement",
    "Changhan Ryu is Software Training Team Leader in the CTO Division at LG Electronics. "
    "He earned an MEd in Learning Design and Technologies from Arizona State University in May 2026. "
    "His academic background also includes computer science at Rutgers University and education at Korea University. "
    "His interests include generative AI–mediated learning, instructional design and workplace learning.",
)
add_field(
    "AI use disclosure",
    "During the preparation of this manuscript, the author used generative AI tools for brainstorming, outlining, language refinement, source mapping, and critical dialogue. The author reviewed, revised, verified, and takes responsibility for all arguments, interpretations, citations, and conclusions. No generative AI tool is listed as an author.",
)

doc.save(OUT)
print("saved", OUT)
