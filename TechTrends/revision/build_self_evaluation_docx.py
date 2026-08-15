"""Build the Korean self-evaluation report as a readable Word document."""
import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


parser = argparse.ArgumentParser()
parser.add_argument("--src", required=True)
parser.add_argument("--out", required=True)
args = parser.parse_args()


def plain(text):
    text = re.sub(r"\[([^]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    return text.replace("**", "").replace("`", "").replace("*", "")


def add_page_field(paragraph):
    run = paragraph.add_run()
    for kind, value in (("begin", None), (None, " PAGE "), ("separate", None), (None, "1"), ("end", None)):
        if kind:
            node = OxmlElement("w:fldChar")
            node.set(qn("w:fldCharType"), kind)
        elif value == " PAGE ":
            node = OxmlElement("w:instrText")
            node.set(qn("xml:space"), "preserve")
            node.text = value
        else:
            node = OxmlElement("w:t")
            node.text = value
        run._r.append(node)


doc = Document()
doc.core_properties.author = ""
doc.core_properties.last_modified_by = ""
doc.core_properties.comments = ""
doc.core_properties.title = "TechTrends Major Revision 자체평가 리포트"
section = doc.sections[0]
for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(section, side, Inches(0.8))

normal = doc.styles["Normal"]
normal.font.name = "Malgun Gothic"
normal.font.size = Pt(10)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.25

for name, size in (("Title", 18), ("Heading 1", 14), ("Heading 2", 12)):
    style = doc.styles[name]
    style.font.name = "Malgun Gothic"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    style.font.size = Pt(size)
    style.font.bold = True

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_page_field(footer)

lines = Path(args.src).read_text(encoding="utf-8").splitlines()
i = 0
first_heading = True
while i < len(lines):
    line = lines[i].rstrip()
    if not line:
        i += 1
        continue
    if line.startswith("| "):
        raw_rows = []
        while i < len(lines) and lines[i].startswith("|"):
            raw_rows.append(lines[i])
            i += 1
        rows = []
        for row in raw_rows:
            cells = [plain(c.strip()) for c in row.strip().strip("|").split("|")]
            if all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
                continue
            rows.append(cells)
        table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
        table.style = "Table Grid"
        for r_idx, cells in enumerate(rows):
            for c_idx, value in enumerate(cells):
                p = table.cell(r_idx, c_idx).paragraphs[0]
                p.add_run(value)
                if r_idx == 0:
                    for run in p.runs:
                        run.bold = True
        doc.add_paragraph()
        continue
    if line.startswith("# "):
        text = plain(line[2:])
        if first_heading:
            p = doc.add_paragraph(text, style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            first_heading = False
        else:
            doc.add_heading(text, level=1)
        i += 1
        continue
    if line.startswith("## "):
        doc.add_heading(plain(line[3:]), level=2)
        i += 1
        continue
    if re.match(r"^\d+\.\s+", line):
        doc.add_paragraph(plain(re.sub(r"^\d+\.\s+", "", line)), style="List Number")
        i += 1
        continue
    if line.startswith("- "):
        doc.add_paragraph(plain(line[2:]), style="List Bullet")
        i += 1
        continue
    doc.add_paragraph(plain(line))
    i += 1

doc.save(args.out)
print("saved", args.out)
