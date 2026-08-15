"""Build a two-column response-to-reviewers Word document."""
import argparse
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/home/ubuntu/work/reverse-learning/TechTrends/revision"
parser = argparse.ArgumentParser()
parser.add_argument("--src", default=f"{BASE}/Response_to_Reviewers.md")
parser.add_argument("--out", default=f"{BASE}/Response_to_Reviewers.docx")
args = parser.parse_args()

with open(args.src) as f:
    text = f.read()

header, *items_raw = text.split("[ITEM]")
items = []
for raw in items_raw:
    source = re.search(r"SOURCE:\s*(.+)", raw).group(1).strip()
    comment = re.search(r"COMMENT:\s*(.+?)\nRESPONSE:", raw, re.S).group(1).strip()
    response = re.search(r"RESPONSE:\s*(.+)", raw, re.S).group(1).strip()
    items.append((source, comment, response))

doc = Document()
doc.core_properties.author = ""
doc.core_properties.last_modified_by = ""
doc.core_properties.comments = ""
doc.core_properties.title = "Response to Reviewers"

section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width, section.page_height = section.page_height, section.page_width
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_node = OxmlElement("w:t")
    text_node.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text_node, end):
        run._r.append(element)


footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_page_field(footer)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(10.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Response to Reviewers")
run.bold = True
run.font.size = Pt(14)

for line in header.strip().split("\n\n")[1:]:
    p = doc.add_paragraph(line.replace("\n", " ").strip())
    p.paragraph_format.space_after = Pt(8)

current_source = None
table = None

def set_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "left", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "666666")
        borders.append(el)
    tblPr.append(borders)

for source, comment, response in items:
    if source != current_source:
        current_source = source
        h = doc.add_heading(source, level=1)
        for r in h.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        set_borders(table)
        hdr = table.rows[0].cells
        trPr = table.rows[0]._tr.get_or_add_trPr()
        repeat = OxmlElement("w:tblHeader")
        repeat.set(qn("w:val"), "true")
        trPr.append(repeat)
        for idx, label in enumerate(("Reviewer suggestion", "Response")):
            para = hdr[idx].paragraphs[0]
            r = para.add_run(label)
            r.bold = True
        for idx, w in enumerate((Inches(4.0), Inches(5.4))):
            for cell in table.columns[idx].cells:
                cell.width = w
    row = table.add_row().cells
    cant_split = OxmlElement("w:cantSplit")
    row[0]._tc.getparent().get_or_add_trPr().append(cant_split)
    row[0].width = Inches(4.0)
    row[1].width = Inches(5.4)
    row[0].paragraphs[0].add_run(comment)
    row[1].paragraphs[0].add_run(response)
    for cell in row:
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(4)

doc.save(args.out)
print("saved", args.out)
